import streamlit as st
import os
import datetime
import pandas as pd
import csv
from utils import (
    convert_mp4_to_wav,
    transcribe_audio,
    transcribe_audio_faster_whisper,
    transcribe_audio_openai_whisper,
    transcribe_audio_whisperx,
    transcribe_with_whisperx_diarization,
    transcribe_with_speakers,
    transcribe_with_speechbrain,
    transcribe_with_resemblyzer,
    extract_speakers_and_updates,
    summarize_updates,
    format_summary_report
)

# Export format imports
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils.dataframe import dataframe_to_rows
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Export Functions
def export_to_word(summary_dict, meeting_date, filename=None):
    """Export meeting summary to Word document"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        title = doc.add_heading(f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        for i, (speaker, summary) in enumerate(summary_dict.items()):
            heading = doc.add_heading(f'👤 {speaker}', level=2)
            doc.add_paragraph(summary)
            if i < len(summary_dict) - 1:
                doc.add_paragraph()
        
        if filename is None:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.docx"
        
        doc.save(filename)
        return filename
    except Exception as e:
        return None

def export_to_pdf(summary_dict, meeting_date, filename=None):
    """Export meeting summary to PDF document"""
    if not PDF_AVAILABLE:
        return None
    
    try:
        if filename is None:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.pdf"
        
        doc = SimpleDocTemplate(filename, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.darkblue,
            alignment=1,
            spaceAfter=30
        )
        
        speaker_style = ParagraphStyle(
            'SpeakerStyle',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.darkgreen,
            spaceAfter=12
        )
        
        title = Paragraph(f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}', title_style)
        story.append(title)
        
        gen_info = Paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal'])
        story.append(gen_info)
        story.append(Spacer(1, 20))
        
        for speaker, summary in summary_dict.items():
            speaker_heading = Paragraph(f'👤 {speaker}', speaker_style)
            story.append(speaker_heading)
            
            content = Paragraph(summary, styles['Normal'])
            story.append(content)
            story.append(Spacer(1, 15))
        
        doc.build(story)
        return filename
    except Exception as e:
        return None

def export_to_excel(summary_dict, meeting_date, filename=None):
    """Export meeting summary to Excel spreadsheet"""
    if not EXCEL_AVAILABLE:
        return None
    
    try:
        if filename is None:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.xlsx"
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Meeting Summary"
        
        ws['A1'] = f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}'
        ws['A1'].font = Font(size=16, bold=True, color="0066CC")
        ws.merge_cells('A1:C1')
        
        ws['A2'] = f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        ws['A2'].font = Font(italic=True)
        ws.merge_cells('A2:C2')
        
        ws['A4'] = 'Speaker'
        ws['B4'] = 'Summary'
        ws['C4'] = 'Word Count'
        
        for cell in ['A4', 'B4', 'C4']:
            ws[cell].font = Font(bold=True, color="FFFFFF")
            ws[cell].fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            ws[cell].alignment = Alignment(horizontal="center")
        
        row = 5
        for speaker, summary in summary_dict.items():
            ws[f'A{row}'] = speaker
            ws[f'B{row}'] = summary
            ws[f'C{row}'] = len(summary.split())
            
            ws[f'A{row}'].font = Font(bold=True)
            ws[f'A{row}'].fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
            
            row += 1
        
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 80
        ws.column_dimensions['C'].width = 15
        
        wb.save(filename)
        return filename
    except Exception as e:
        return None

def export_to_csv(summary_dict, meeting_date, filename=None):
    """Export meeting summary to CSV file"""
    try:
        if filename is None:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.csv"
        
        data = []
        data.append(['Meeting Summary', meeting_date.strftime("%Y-%m-%d")])
        data.append(['Generated on', datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        data.append([])
        data.append(['Speaker', 'Summary', 'Word Count'])
        
        for speaker, summary in summary_dict.items():
            word_count = len(summary.split())
            data.append([speaker, summary, word_count])
        
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerows(data)
        
        return filename
    except Exception as e:
        return None

def check_export_dependencies():
    """Check which export formats are available"""
    available_formats = ['csv']
    
    if DOCX_AVAILABLE:
        available_formats.append('word')
    if PDF_AVAILABLE:
        available_formats.append('pdf')
    if EXCEL_AVAILABLE:
        available_formats.append('excel')
    
    return available_formats

st.set_page_config(page_title="Meeting Progress Summarizer", layout="wide")

# Custom CSS for better button styling
st.markdown("""
<style>
    /* Make buttons more responsive and consistent */
    .stButton > button {
        height: 2.5rem;
        border-radius: 8px;
        font-weight: 500;
        transition: all 0.2s ease;
    }
    
    /* Primary button styling */
    .stButton > button[data-baseweb="button"][data-testid="baseButton-primary"] {
        background-color: #0066cc;
        border: 2px solid #0066cc;
        box-shadow: 0 2px 4px rgba(0,102,204,0.3);
    }
    
    /* Secondary button styling */
    .stButton > button[data-baseweb="button"][data-testid="baseButton-secondary"] {
        background-color: #f0f2f6;
        border: 2px solid #e0e4ea;
        color: #262730;
    }
    
    /* Hover effects */
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    
    /* Disabled button styling */
    .stButton > button:disabled {
        opacity: 0.4;
        cursor: not-allowed;
    }
    
    /* Configuration section styling */
    .stSubheader {
        color: #0066cc;
        border-bottom: 2px solid #e0e4ea;
        padding-bottom: 0.5rem;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

st.title("📋 Meeting Transcript Summarizer")

uploaded_file = st.file_uploader("Upload MP4 meeting recording", type=["mp4"])

# Configuration section with columns
st.subheader("🔧 Configuration")

# Initialize session state for selections
if 'transcription_engine' not in st.session_state:
    st.session_state.transcription_engine = 'faster-whisper'
if 'selected_model' not in st.session_state:
    st.session_state.selected_model = 'small.en'
if 'speaker_method' not in st.session_state:
    st.session_state.speaker_method = 'llm-only'
if 'mode' not in st.session_state:
    st.session_state.mode = 'LLM'
if 'llm_model' not in st.session_state:
    st.session_state.llm_model = 'qwen2:latest'

col1, col2, col3 = st.columns(3)

with col1:
    st.write("**Transcription Engine:**")
    
    # Transcription Engine Buttons
    engine_col1, engine_col2 = st.columns(2)
    with engine_col1:
        if st.button("🚀 faster-whisper", 
                    type="primary" if st.session_state.transcription_engine == 'faster-whisper' else "secondary",
                    use_container_width=True,
                    key="btn_faster_whisper"):
            st.session_state.transcription_engine = 'faster-whisper'
    
    with engine_col2:
        if st.button("🔄 openai-whisper", 
                    type="primary" if st.session_state.transcription_engine == 'openai-whisper' else "secondary",
                    use_container_width=True,
                    key="btn_openai_whisper"):
            st.session_state.transcription_engine = 'openai-whisper'
    
    st.caption(f"Selected: **{st.session_state.transcription_engine}**")

with col2:
    st.write("**Whisper Model:**")
    
    # Model info helper
    model_info = {
        "tiny.en": "39 MB, ~32x faster, English only",
        "tiny": "39 MB, ~32x faster, multilingual", 
        "base.en": "74 MB, ~16x faster, English only",
        "base": "74 MB, ~16x faster, multilingual",
        "small.en": "244 MB, ~6x faster, English only",
        "small": "244 MB, ~6x faster, multilingual",
        "medium.en": "769 MB, ~2x faster, English only", 
        "medium": "769 MB, ~2x faster, multilingual",
        "large-v1": "1550 MB, 1x speed, multilingual",
        "large-v2": "1550 MB, 1x speed, multilingual", 
        "large-v3": "1550 MB, 1x speed, multilingual",
        "large": "1550 MB, 1x speed, multilingual"
    }
    
    # Whisper Model Buttons (organized in rows)
    model_col1, model_col2 = st.columns(2)
    
    models_row1 = ["tiny.en", "tiny", "base.en", "base"]
    models_row2 = ["small.en", "small", "medium.en", "medium"]
    models_row3 = ["large-v1", "large-v2", "large-v3", "large"]
    
    # Row 1
    for i, model in enumerate(models_row1):
        col = model_col1 if i % 2 == 0 else model_col2
        with col:
            if st.button(f"📏 {model}", 
                        type="primary" if st.session_state.selected_model == model else "secondary",
                        use_container_width=True,
                        key=f"btn_model_{model}"):
                st.session_state.selected_model = model
    
    # Row 2
    for i, model in enumerate(models_row2):
        col = model_col1 if i % 2 == 0 else model_col2
        with col:
            if st.button(f"📏 {model}", 
                        type="primary" if st.session_state.selected_model == model else "secondary",
                        use_container_width=True,
                        key=f"btn_model_{model}"):
                st.session_state.selected_model = model
    
    # Row 3
    for i, model in enumerate(models_row3):
        col = model_col1 if i % 2 == 0 else model_col2
        with col:
            if st.button(f"📏 {model}", 
                        type="primary" if st.session_state.selected_model == model else "secondary",
                        use_container_width=True,
                        key=f"btn_model_{model}"):
                st.session_state.selected_model = model
    
    st.caption(f"ℹ️ {model_info.get(st.session_state.selected_model, 'Model info not available')}")

with col3:
    st.write("**Speaker Detection Method:**")
    
    # Speaker Detection Method Buttons
    speaker_methods = ["llm-only", "pyannote", "speechbrain", "resemblyzer"]
    speaker_labels = ["🤖 LLM Only", "🎯 Pyannote", "🧠 SpeechBrain", "🔊 Resemblyzer"]
    
    for method, label in zip(speaker_methods, speaker_labels):
        if st.button(label, 
                    type="primary" if st.session_state.speaker_method == method else "secondary",
                    use_container_width=True,
                    key=f"btn_speaker_{method}"):
            st.session_state.speaker_method = method
    
    st.caption(f"Selected: **{st.session_state.speaker_method}**")

# Additional LLM Configuration (below the main columns)
st.write("**LLM Configuration:**")
llm_col1, llm_col2 = st.columns(2)

with llm_col1:
    if st.session_state.speaker_method == "llm-only":
        st.write("*LLM Detection Mode:*")
        mode_col1, mode_col2, mode_col3 = st.columns(3)
        
        with mode_col1:
            if st.button("🤖 LLM", 
                        type="primary" if st.session_state.mode == 'LLM' else "secondary",
                        use_container_width=True,
                        key="btn_mode_llm"):
                st.session_state.mode = 'LLM'
        
        with mode_col2:
            if st.button("📋 Strict", 
                        type="primary" if st.session_state.mode == 'Strict' else "secondary",
                        use_container_width=True,
                        key="btn_mode_strict"):
                st.session_state.mode = 'Strict'
        
        with mode_col3:
            if st.button("🧠 Smart", 
                        type="primary" if st.session_state.mode == 'Smart' else "secondary",
                        use_container_width=True,
                        key="btn_mode_smart"):
                st.session_state.mode = 'Smart'

with llm_col2:
    st.write("*LLM Model:*")
    llm_models = ["qwen2:latest", "phi3:latest", "llama2:latest", "gemma:7b", "mistral:latest"]
    llm_labels = ["🔥 Qwen2", "⚡ Phi3", "🦙 Llama2", "💎 Gemma", "🌟 Mistral"]
    
    # Organize LLM buttons in 2 columns
    llm_btn_col1, llm_btn_col2 = st.columns(2)
    
    for i, (model, label) in enumerate(zip(llm_models, llm_labels)):
        col = llm_btn_col1 if i % 2 == 0 else llm_btn_col2
        with col:
            if st.button(label, 
                        type="primary" if st.session_state.llm_model == model else "secondary",
                        use_container_width=True,
                        key=f"btn_llm_{model.replace(':', '_').replace('.', '_')}"):
                st.session_state.llm_model = model

# Store the selected values for use in the processing
transcription_engine = st.session_state.transcription_engine
selected_model = st.session_state.selected_model
speaker_method = st.session_state.speaker_method
mode = st.session_state.mode
llm_model = st.session_state.llm_model

# Sidebar with model information and tips
st.sidebar.header("📊 Model Information")
st.sidebar.markdown("""
### 🚀 Transcription Engines:
- **faster-whisper**: Optimized, uses less memory
- **openai-whisper**: Original implementation with full features

### 📏 Model Sizes & Performance:
- **tiny**: Fastest, least accurate
- **small**: Good balance of speed/accuracy  
- **medium**: Better accuracy, slower
- **large**: Best accuracy, slowest

### 💡 Tips:
- Use **English-only** models (.en) for English audio
- **small.en** recommended for most use cases
- **medium** or **large** for critical accuracy
- **whisperx** provides best speaker separation
""")

st.sidebar.markdown("---")
st.sidebar.markdown("### 🎯 Speaker Detection Methods:")
st.sidebar.markdown("""
- **llm-only**: Pure AI text analysis
- **pyannote**: Advanced audio-based diarization  
- **speechbrain**: Speaker recognition with embeddings
- **resemblyzer**: Voice verification clustering
""")

st.sidebar.markdown("---")
st.sidebar.markdown("### 🤖 Available LLM Models:")
st.sidebar.markdown("""
- **qwen2:latest**: Alibaba Qwen2 latest, excellent reasoning and multilingual capabilities
- **phi3:latest**: Microsoft Phi-3 latest, efficient and powerful for complex tasks
- **llama2:latest**: Meta Llama 2 latest, proven general-purpose model
- **gemma:7b**: Google Gemma 7B, excellent balance of performance and efficiency
- **mistral:latest**: Latest Mistral model with strong performance

### 📋 Model Recommendations:
- **Best reasoning**: Use **qwen2:latest** or **phi3:latest**
- **Proven performance**: Use **llama2:latest** or **mistral:latest**
- **Balanced efficiency**: Use **gemma:7b**
- **Complex analysis**: Use **qwen2:latest**
- **General purpose**: Use **llama2:latest**
""")

st.sidebar.markdown("---")
st.sidebar.markdown("### 🔧 LLM Detection Modes:")
st.sidebar.markdown("""
- **LLM**: AI-powered flexible detection
- **Strict**: Pattern-based ("Hi Team, this is...")  
- **Smart**: NLP entity recognition
""")

meeting_date = st.date_input("Meeting Date (optional)", value=datetime.date.today())

if uploaded_file:
    with st.spinner("🔄 Converting video to audio..."):
        wav_path = convert_mp4_to_wav(uploaded_file)

    # Display selected configuration
    st.success(f"✅ Configuration: {transcription_engine} with {selected_model} model, {speaker_method} speaker detection")
    
    # Step 1: Transcription and Speaker Detection
    if speaker_method == "pyannote":
        # Use Pyannote for speaker diarization
        with st.spinner(f"🗣️ Transcribing with {transcription_engine} ({selected_model})..."):
            transcript = transcribe_audio(wav_path, transcription_engine, selected_model)
        st.success("✅ Transcription complete")
        
        with st.spinner("🎯 Detecting speakers with Pyannote..."):
            speaker_updates = transcribe_with_speakers(wav_path)
            
    elif speaker_method == "speechbrain":
        # Use SpeechBrain for speaker recognition  
        with st.spinner(f"🗣️ Transcribing with {transcription_engine} ({selected_model})..."):
            transcript = transcribe_audio(wav_path, transcription_engine, selected_model)
        st.success("✅ Transcription complete")
        
        with st.spinner("🎯 Detecting speakers with SpeechBrain..."):
            speaker_updates = transcribe_with_speechbrain(wav_path)
            
    elif speaker_method == "resemblyzer":
        # Use Resemblyzer for speaker verification
        with st.spinner(f"�️ Transcribing with {transcription_engine} ({selected_model})..."):
            transcript = transcribe_audio(wav_path, transcription_engine, selected_model)
        st.success("✅ Transcription complete")
        
        with st.spinner("🎯 Detecting speakers with Resemblyzer..."):
            speaker_updates = transcribe_with_resemblyzer(wav_path)
            
    else:  # llm-only
        # Traditional LLM-based speaker extraction
        with st.spinner(f"🗣️ Transcribing with {transcription_engine} ({selected_model})..."):
            transcript = transcribe_audio(wav_path, transcription_engine, selected_model)
        st.success("✅ Transcription complete")
        
        with st.spinner("🧑‍💼 Extracting speakers with LLM..."):
            speaker_updates = extract_speakers_and_updates(transcript, mode, llm_model)

    # Show transcript preview with expandable section
    with st.expander("📝 View Transcript Preview"):
        st.text_area("Transcript", transcript, height=200)

    # Show speaker breakdown
    if speaker_updates:
        st.success(f"✅ Detected {len(speaker_updates)} speakers")
        
        with st.expander("👥 View Speaker Breakdown"):
            for speaker, content in speaker_updates.items():
                st.subheader(f"🎤 {speaker}")
                st.write(content[:200] + "..." if len(content) > 200 else content)
                st.divider()

        with st.spinner("🧠 Summarizing each speaker's progress..."):
            summaries = summarize_updates(speaker_updates, llm_model)

        with st.spinner("📄 Generating summary report..."):
            report = format_summary_report(summaries, meeting_date)

        # Display results
        st.subheader("📋 Meeting Summary")
        
        # Show individual summaries
        for speaker, summary in summaries.items():
            with st.container():
                st.markdown(f"**👤 {speaker}:**")
                st.write(summary)
                st.divider()

        # Export Options Section
        st.subheader("📤 Export Options")
        
        # Check available export formats
        available_formats = check_export_dependencies()
        
        # Initialize session state for export formats
        if 'export_formats' not in st.session_state:
            st.session_state.export_formats = ['csv']  # Default to CSV
        
        # Display available formats and selection buttons
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.write("**Select export formats:**")
            
            # Export format buttons
            format_col1, format_col2 = st.columns(2)
            
            format_info = {
                'csv': ('📊 CSV', 'Simple spreadsheet format'),
                'word': ('📝 Word', 'Professional document'),
                'pdf': ('📄 PDF', 'Universal document format'),
                'excel': ('📈 Excel', 'Advanced spreadsheet')
            }
            
            for i, fmt in enumerate(['csv', 'word', 'pdf', 'excel']):
                col = format_col1 if i % 2 == 0 else format_col2
                with col:
                    is_available = fmt in available_formats
                    is_selected = fmt in st.session_state.export_formats
                    
                    button_type = "primary" if is_selected else "secondary"
                    button_disabled = not is_available
                    
                    label, desc = format_info[fmt]
                    
                    if st.button(
                        f"{label}", 
                        type=button_type,
                        disabled=button_disabled,
                        use_container_width=True,
                        key=f"btn_export_{fmt}",
                        help=desc if is_available else f"Install dependencies for {fmt.upper()}"
                    ):
                        if is_selected:
                            # Remove from selection
                            if fmt in st.session_state.export_formats:
                                st.session_state.export_formats.remove(fmt)
                        else:
                            # Add to selection
                            if fmt not in st.session_state.export_formats:
                                st.session_state.export_formats.append(fmt)
            
            # Show selected formats
            if st.session_state.export_formats:
                selected_text = ", ".join([f.upper() for f in st.session_state.export_formats])
                st.success(f"🎯 Selected: **{selected_text}**")
            else:
                st.warning("⚠️ No export formats selected")
        
        with col2:
            st.write("**Format Status:**")
            format_status = {
                'csv': '✅ Always available',
                'word': '✅ Available' if DOCX_AVAILABLE else '❌ Install python-docx',
                'pdf': '✅ Available' if PDF_AVAILABLE else '❌ Install reportlab', 
                'excel': '✅ Available' if EXCEL_AVAILABLE else '❌ Install openpyxl'
            }
            
            for fmt in ['csv', 'word', 'pdf', 'excel']:
                st.caption(f"{fmt.upper()}: {format_status[fmt]}")

        # Export button
        if st.session_state.export_formats:
            st.write("")  # Add some space
            
            if st.button("� Generate & Export Files", type="primary", use_container_width=True):
                export_results = {}
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                
                with st.spinner("📤 Exporting files..."):
                    for fmt in st.session_state.export_formats:
                        if fmt == 'word' and DOCX_AVAILABLE:
                            filename = f"meeting_summary_{timestamp}.docx"
                            result = export_to_word(summaries, meeting_date, filename)
                        elif fmt == 'pdf' and PDF_AVAILABLE:
                            filename = f"meeting_summary_{timestamp}.pdf"
                            result = export_to_pdf(summaries, meeting_date, filename)
                        elif fmt == 'excel' and EXCEL_AVAILABLE:
                            filename = f"meeting_summary_{timestamp}.xlsx"
                            result = export_to_excel(summaries, meeting_date, filename)
                        elif fmt == 'csv':
                            filename = f"meeting_summary_{timestamp}.csv"
                            result = export_to_csv(summaries, meeting_date, filename)
                        else:
                            result = None
                        
                        export_results[fmt] = result
                
                # Show export results and download buttons
                successful_exports = {fmt: path for fmt, path in export_results.items() if path}
                failed_exports = {fmt: path for fmt, path in export_results.items() if not path}
                
                if successful_exports:
                    st.success(f"✅ Successfully exported {len(successful_exports)} file(s)")
                    
                    # Create download buttons for each successful export
                    download_col1, download_col2 = st.columns(2)
                    
                    for i, (fmt, filepath) in enumerate(successful_exports.items()):
                        col = download_col1 if i % 2 == 0 else download_col2
                        
                        if filepath and os.path.exists(filepath):
                            with open(filepath, 'rb') as file:
                                file_data = file.read()
                                
                            # Determine MIME type
                            mime_types = {
                                'word': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                'pdf': 'application/pdf',
                                'excel': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                'csv': 'text/csv'
                            }
                            
                            filename = os.path.basename(filepath)
                            
                            with col:
                                st.download_button(
                                    label=f"� {fmt.upper()}",
                                    data=file_data,
                                    file_name=filename,
                                    mime=mime_types.get(fmt, 'application/octet-stream'),
                                    key=f"download_{fmt}_{timestamp}",
                                    use_container_width=True
                                )
                    
                    # Cleanup files after showing download buttons
                    for filepath in successful_exports.values():
                        try:
                            if filepath and os.path.exists(filepath):
                                os.unlink(filepath)
                        except:
                            pass
                
                if failed_exports:
                    st.error(f"❌ Failed to export: {', '.join(failed_exports.keys())}")

        # Original TXT download (keep existing functionality)
        st.divider()
        st.download_button(
            "📥 Download Complete Summary (TXT)", 
            data=report, 
            file_name=f"meeting_summary_{meeting_date.strftime('%Y%m%d')}.txt", 
            mime="text/plain"
        )
        
        with st.expander("📄 View Complete Report"):
            st.text_area("Full Meeting Summary Report", report, height=400)
    else:
        st.error("❌ No speakers detected. Try adjusting the speaker detection mode.")

