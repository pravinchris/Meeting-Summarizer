#!/usr/bin/env python3
"""
Enhanced meeting summarizer with export functionality
"""

import streamlit as st
from datetime import datetime
import os
import tempfile
import pandas as pd
import csv
from utils import (
    convert_mp4_to_wav, 
    transcribe_audio, 
    extract_with_llm, 
    summarize_updates
)

# Export format imports
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils.dataframe import dataframe_to_rows
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Export Functions
def export_to_word(summary_dict, meeting_date, filename=None):
    """Export meeting summary to Word document"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        title = doc.add_heading(f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        for i, (speaker, summary) in enumerate(summary_dict.items()):
            heading = doc.add_heading(f'👤 {speaker}', level=2)
            doc.add_paragraph(summary)
            if i < len(summary_dict) - 1:
                doc.add_paragraph()
        
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.docx"
        
        doc.save(filename)
        return filename
    except Exception as e:
        return None

def export_to_pdf(summary_dict, meeting_date, filename=None):
    """Export meeting summary to PDF document"""
    if not PDF_AVAILABLE:
        return None
    
    try:
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.pdf"
        
        doc = SimpleDocTemplate(filename, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.darkblue,
            alignment=1,
            spaceAfter=30
        )
        
        speaker_style = ParagraphStyle(
            'SpeakerStyle',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.darkgreen,
            spaceAfter=12
        )
        
        title = Paragraph(f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}', title_style)
        story.append(title)
        
        gen_info = Paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal'])
        story.append(gen_info)
        story.append(Spacer(1, 20))
        
        for speaker, summary in summary_dict.items():
            speaker_heading = Paragraph(f'👤 {speaker}', speaker_style)
            story.append(speaker_heading)
            
            content = Paragraph(summary, styles['Normal'])
            story.append(content)
            story.append(Spacer(1, 15))
        
        doc.build(story)
        return filename
    except Exception as e:
        return None

def export_to_excel(summary_dict, meeting_date, filename=None):
    """Export meeting summary to Excel spreadsheet"""
    if not EXCEL_AVAILABLE:
        return None
    
    try:
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.xlsx"
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Meeting Summary"
        
        ws['A1'] = f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}'
        ws['A1'].font = Font(size=16, bold=True, color="0066CC")
        ws.merge_cells('A1:C1')
        
        ws['A2'] = f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        ws['A2'].font = Font(italic=True)
        ws.merge_cells('A2:C2')
        
        ws['A4'] = 'Speaker'
        ws['B4'] = 'Summary'
        ws['C4'] = 'Word Count'
        
        for cell in ['A4', 'B4', 'C4']:
            ws[cell].font = Font(bold=True, color="FFFFFF")
            ws[cell].fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            ws[cell].alignment = Alignment(horizontal="center")
        
        row = 5
        for speaker, summary in summary_dict.items():
            ws[f'A{row}'] = speaker
            ws[f'B{row}'] = summary
            ws[f'C{row}'] = len(summary.split())
            
            ws[f'A{row}'].font = Font(bold=True)
            ws[f'A{row}'].fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
            
            row += 1
        
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 80
        ws.column_dimensions['C'].width = 15
        
        wb.save(filename)
        return filename
    except Exception as e:
        return None

def export_to_csv(summary_dict, meeting_date, filename=None):
    """Export meeting summary to CSV file"""
    try:
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.csv"
        
        data = []
        data.append(['Meeting Summary', meeting_date.strftime("%Y-%m-%d")])
        data.append(['Generated on', datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        data.append([])
        data.append(['Speaker', 'Summary', 'Word Count'])
        
        for speaker, summary in summary_dict.items():
            word_count = len(summary.split())
            data.append([speaker, summary, word_count])
        
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerows(data)
        
        return filename
    except Exception as e:
        return None

def export_summary(summary_dict, meeting_date, formats=['word', 'pdf', 'excel', 'csv']):
    """Export meeting summary to multiple formats"""
    results = {}
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    for format_type in formats:
        format_type = format_type.lower()
        
        if format_type == 'word' and DOCX_AVAILABLE:
            filename = f"meeting_summary_{timestamp}.docx"
            result = export_to_word(summary_dict, meeting_date, filename)
        elif format_type == 'pdf' and PDF_AVAILABLE:
            filename = f"meeting_summary_{timestamp}.pdf"
            result = export_to_pdf(summary_dict, meeting_date, filename)
        elif format_type == 'excel' and EXCEL_AVAILABLE:
            filename = f"meeting_summary_{timestamp}.xlsx"
            result = export_to_excel(summary_dict, meeting_date, filename)
        elif format_type == 'csv':
            filename = f"meeting_summary_{timestamp}.csv"
            result = export_to_csv(summary_dict, meeting_date, filename)
        else:
            result = None
        
        results[format_type] = result
    
    return results

def check_export_dependencies():
    """Check which export formats are available"""
    available_formats = ['csv']
    
    if DOCX_AVAILABLE:
        available_formats.append('word')
    if PDF_AVAILABLE:
        available_formats.append('pdf')
    if EXCEL_AVAILABLE:
        available_formats.append('excel')
    
    return available_formats

def main():
    st.title("🎙️ Meeting Summarizer with Export")
    st.markdown("Upload an audio/video file to generate speaker summaries and export them in various formats")
    
    # Check export capabilities
    with st.expander("📋 Export Format Availability", expanded=False):
        available_formats = check_export_dependencies()
        if len(available_formats) == 4:
            st.success("✅ All export formats available!")
        else:
            st.warning(f"⚠️ {len(available_formats)}/4 export formats available")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose an audio or video file", 
        type=['mp4', 'wav', 'mp3', 'm4a', 'webm']
    )
    
    if uploaded_file is not None:
        st.success(f"📁 File uploaded: {uploaded_file.name}")
        
        # Configuration
        col1, col2, col3 = st.columns(3)
        
        with col1:
            engine = st.selectbox(
                "Transcription Engine",
                ["faster-whisper", "openai-whisper", "whisperx"],
                index=0
            )
        
        with col2:
            model = st.selectbox(
                "Model Size",
                ["tiny.en", "small.en", "base.en", "medium.en"],
                index=1
            )
        
        with col3:
            llm_model = st.selectbox(
                "LLM Model",
                ["llama3.2:1b", "llama3.2:3b", "llama3.1:8b"],
                index=0
            )
        
        # Export format selection
        st.subheader("📤 Export Options")
        export_formats = st.multiselect(
            "Select export formats",
            available_formats,
            default=available_formats[:2] if len(available_formats) >= 2 else available_formats
        )
        
        # Process button
        if st.button("🚀 Process Meeting", type="primary"):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                # Step 1: Convert audio if needed
                status_text.text("🔄 Converting audio...")
                progress_bar.progress(10)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                    if uploaded_file.name.endswith('.wav'):
                        tmp_file.write(uploaded_file.read())
                        wav_path = tmp_file.name
                    else:
                        wav_path = convert_mp4_to_wav(uploaded_file)
                
                # Step 2: Transcribe
                status_text.text("🎯 Transcribing audio...")
                progress_bar.progress(30)
                
                transcript = transcribe_audio(wav_path, engine=engine, model_name=model)
                
                # Step 3: Extract speakers
                status_text.text("👥 Extracting speakers...")
                progress_bar.progress(60)
                
                speakers = extract_with_llm(transcript, llm_model)
                
                # Step 4: Summarize
                status_text.text("📝 Generating summaries...")
                progress_bar.progress(80)
                
                summaries = summarize_updates(speakers, llm_model)
                
                # Step 5: Export
                if export_formats:
                    status_text.text("📤 Exporting files...")
                    progress_bar.progress(90)
                    
                    meeting_date = datetime.now()
                    export_results = export_summary(
                        summary_dict=summaries,
                        meeting_date=meeting_date,
                        formats=export_formats
                    )
                else:
                    export_results = {}
                
                progress_bar.progress(100)
                status_text.text("✅ Processing complete!")
                
                # Display results
                st.success("🎉 Meeting summary generated successfully!")
                
                # Show transcript preview
                with st.expander("📜 Full Transcript", expanded=False):
                    st.text_area("Transcript", transcript, height=200)
                
                # Show speaker summaries
                st.subheader("👥 Speaker Summaries")
                for speaker, summary in summaries.items():
                    with st.container():
                        st.markdown(f"**👤 {speaker}**")
                        st.write(summary)
                        st.divider()
                
                # Show export results
                if export_results:
                    st.subheader("📁 Exported Files")
                    
                    successful_exports = {fmt: path for fmt, path in export_results.items() if path}
                    failed_exports = {fmt: path for fmt, path in export_results.items() if not path}
                    
                    if successful_exports:
                        st.success(f"✅ Successfully exported {len(successful_exports)} file(s):")
                        for fmt, filepath in successful_exports.items():
                            if filepath and os.path.exists(filepath):
                                with open(filepath, 'rb') as file:
                                    file_data = file.read()
                                    
                                # Determine MIME type
                                mime_types = {
                                    'word': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                    'pdf': 'application/pdf',
                                    'excel': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                    'csv': 'text/csv'
                                }
                                
                                filename = os.path.basename(filepath)
                                st.download_button(
                                    label=f"📄 Download {fmt.upper()} - {filename}",
                                    data=file_data,
                                    file_name=filename,
                                    mime=mime_types.get(fmt, 'application/octet-stream'),
                                    key=f"download_{fmt}"
                                )
                    
                    if failed_exports:
                        st.error(f"❌ Failed to export {len(failed_exports)} file(s): {', '.join(failed_exports.keys())}")
                
                # Cleanup
                try:
                    os.unlink(wav_path)
                    for filepath in export_results.values():
                        if filepath and os.path.exists(filepath):
                            os.unlink(filepath)
                except:
                    pass
                
            except Exception as e:
                st.error(f"❌ Error processing meeting: {str(e)}")
                progress_bar.empty()
                status_text.empty()

if __name__ == "__main__":
    main()
