import os
import subprocess
from datetime import datetime
from faster_whisper import WhisperModel
import tempfile
import ollama
import warnings
import pandas as pd
import csv

# Export format imports
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. Install with: pip install python-docx")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("reportlab not available. Install with: pip install reportlab")

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils.dataframe import dataframe_to_rows
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("openpyxl not available. Install with: pip install openpyxl")

# Suppress warnings for clean output
warnings.filterwarnings("ignore")

try:
    import whisper
    OPENAI_WHISPER_AVAILABLE = True
except ImportError:
    OPENAI_WHISPER_AVAILABLE = False
    print("OpenAI Whisper not available. Install with: pip install whisper")

try:
    import whisperx
    WHISPERX_AVAILABLE = True
except ImportError:
    WHISPERX_AVAILABLE = False
    print("WhisperX not available. Install with: pip install whisperx")
except Exception as e:
    WHISPERX_AVAILABLE = False
    print(f"WhisperX not available due to dependency issues: {e}")

# Convert MP4 to WAV using ffmpeg
def convert_mp4_to_wav(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp_mp4:
        tmp_mp4.write(uploaded_file.read())
        tmp_mp4.flush()
        wav_path = tmp_mp4.name.replace(".mp4", ".wav")
        subprocess.run(["ffmpeg", "-i", tmp_mp4.name, "-ar", "16000", "-ac", "1", wav_path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return wav_path

# Transcribe using faster-whisper with model selection
def transcribe_audio_faster_whisper(wav_path, model_name="small.en"):
    """
    Transcribe audio using faster-whisper with specified model
    Available models: tiny.en, tiny, base.en, base, small.en, small, medium.en, medium, large-v1, large-v2, large-v3, large
    """
    try:
        model = WhisperModel(model_name, compute_type="auto")
        segments, _ = model.transcribe(wav_path)
        return " ".join([seg.text.strip() for seg in segments])
    except Exception as e:
        print(f"Error with faster-whisper model {model_name}: {e}")
        # Fallback to smaller model
        fallback_model = "tiny.en" if model_name != "tiny.en" else "tiny"
        print(f"Falling back to {fallback_model}")
        model = WhisperModel(fallback_model, compute_type="auto")
        segments, _ = model.transcribe(wav_path)
        return " ".join([seg.text.strip() for seg in segments])

# Transcribe using OpenAI Whisper with model selection  
def transcribe_audio_openai_whisper(wav_path, model_name="small.en"):
    """
    Transcribe audio using OpenAI Whisper with specified model
    Available models: tiny.en, tiny, base.en, base, small.en, small, medium.en, medium, large-v1, large-v2, large-v3, large
    """
    if not OPENAI_WHISPER_AVAILABLE:
        print("OpenAI Whisper not available, falling back to faster-whisper")
        return transcribe_audio_faster_whisper(wav_path, model_name)
    
    try:
        model = whisper.load_model(model_name)
        result = model.transcribe(wav_path)
        return result["text"].strip()
    except Exception as e:
        print(f"Error with OpenAI Whisper model {model_name}: {e}")
        # Fallback to smaller model
        fallback_model = "tiny.en" if model_name != "tiny.en" else "tiny"
        print(f"Falling back to {fallback_model}")
        try:
            model = whisper.load_model(fallback_model)
            result = model.transcribe(wav_path)
            return result["text"].strip()
        except Exception as fallback_e:
            print(f"Fallback also failed: {fallback_e}, using faster-whisper instead")
            return transcribe_audio_faster_whisper(wav_path, fallback_model)

# Transcribe using WhisperX with model selection and enhanced features
def transcribe_audio_whisperx(wav_path, model_name="small.en", device="cpu", batch_size=16):
    """
    Transcribe audio using WhisperX with specified model
    WhisperX provides better timestamp accuracy and word-level alignment
    Available models: tiny.en, tiny, base.en, base, small.en, small, medium.en, medium, large-v1, large-v2, large-v3, large
    """
    if not WHISPERX_AVAILABLE:
        print("WhisperX not available, falling back to faster-whisper")
        return transcribe_audio_faster_whisper(wav_path, model_name)
    
    try:
        import torch
        
        # Auto-detect device if not specified
        if device == "auto":
            device = "cuda" if torch.cuda.is_available() else "cpu"
        
        compute_type = "float16" if device == "cuda" else "int8"
        
        # Load WhisperX model
        model = whisperx.load_model(model_name, device, compute_type=compute_type)
        
        # Load audio
        audio = whisperx.load_audio(wav_path)
        
        # Perform transcription
        result = model.transcribe(audio, batch_size=batch_size)
        
        # Extract text from segments
        if "segments" in result:
            transcript = " ".join([segment["text"].strip() for segment in result["segments"]])
        else:
            transcript = result.get("text", "").strip()
        
        return transcript
        
    except Exception as e:
        print(f"Error with WhisperX model {model_name}: {e}")
        # Fallback to faster-whisper
        print("Falling back to faster-whisper")
        return transcribe_audio_faster_whisper(wav_path, model_name)

# Transcribe using WhisperX with speaker diarization
def transcribe_with_whisperx_diarization(wav_path, model_name="small.en", device="cpu", batch_size=16):
    """
    Transcribe audio using WhisperX with built-in speaker diarization
    Returns a dictionary with speaker labels and their corresponding text
    """
    if not WHISPERX_AVAILABLE:
        print("WhisperX not available, falling back to pyannote")
        return transcribe_with_speakers(wav_path)
    
    try:
        import torch
        
        # Auto-detect device if not specified
        if device == "auto":
            device = "cuda" if torch.cuda.is_available() else "cpu"
        
        compute_type = "float16" if device == "cuda" else "int8"
        
        # Load WhisperX model
        model = whisperx.load_model(model_name, device, compute_type=compute_type)
        
        # Load audio
        audio = whisperx.load_audio(wav_path)
        
        # Perform transcription
        result = model.transcribe(audio, batch_size=batch_size)
        
        # Load alignment model for better timestamps
        try:
            model_a, metadata = whisperx.load_align_model(language_code=result["language"], device=device)
            result = whisperx.align(result["segments"], model_a, metadata, audio, device, return_char_alignments=False)
        except Exception as align_e:
            print(f"Alignment failed: {align_e}, continuing without alignment")
        
        # Perform speaker diarization
        try:
            # Load diarization model
            diarize_model = whisperx.DiarizationPipeline(use_auth_token=None, device=device)
            
            # Perform diarization
            diarize_segments = diarize_model(audio)
            
            # Assign speakers to segments
            result = whisperx.assign_word_speakers(diarize_segments, result)
            
            # Group by speakers
            speaker_updates = {}
            for segment in result["segments"]:
                speaker = segment.get("speaker", "Unknown_Speaker")
                text = segment["text"].strip()
                
                if speaker in speaker_updates:
                    speaker_updates[speaker] += " " + text
                else:
                    speaker_updates[speaker] = text
            
            return speaker_updates if speaker_updates else {"Speaker_1": " ".join([seg["text"] for seg in result["segments"]])}
            
        except Exception as diarize_e:
            print(f"Diarization failed: {diarize_e}, returning transcript without speakers")
            # Return single speaker if diarization fails
            transcript = " ".join([segment["text"].strip() for segment in result["segments"]])
            return {"Speaker_1": transcript}
        
    except Exception as e:
        print(f"Error with WhisperX diarization: {e}")
        # Fallback to pyannote
        print("Falling back to pyannote speaker diarization")
        return transcribe_with_speakers(wav_path)

# Updated main transcription function with WhisperX support
def transcribe_audio(wav_path, engine="faster-whisper", model_name="small.en", device="auto", batch_size=16):
    """
    Main transcription function that supports all three engines
    """
    if engine == "whisperx":
        return transcribe_audio_whisperx(wav_path, model_name, device, batch_size)
    elif engine == "faster-whisper":
        return transcribe_audio_faster_whisper(wav_path, model_name)
    elif engine == "openai-whisper":
        return transcribe_audio_openai_whisper(wav_path, model_name)
    else:
        # Default fallback
        return transcribe_audio_faster_whisper(wav_path, model_name)

# LLM prompt to extract speaker updates
def extract_speakers_and_updates(transcript, mode, model_name):
    if mode == "Strict":
        return extract_strict(transcript)
    elif mode == "Smart":
        return extract_smart(transcript)
    else:
        return extract_with_llm(transcript, model_name)

def extract_with_llm(transcript, model_name):
    """
    Extract speakers using LLM with improved prompting and parsing to identify actual names
    """
    prompt = f"""Analyze this meeting transcript and identify the actual names of speakers and what they said. 

TRANSCRIPT:
{transcript}

INSTRUCTIONS:
1. Look for speaker introductions like "Hi, I'm John", "This is Sarah", "My name is Mike", etc.
2. Identify actual names mentioned in the conversation
3. Extract what each person contributed to the discussion
4. If you can't find real names, try to identify speakers by context clues
5. Return results in this exact format:

[ACTUAL_NAME_1]: [their contribution]
[ACTUAL_NAME_2]: [their contribution]
[ACTUAL_NAME_3]: [their contribution]

EXAMPLES:
- If someone says "Hi, I'm Sarah and I worked on the database", use "Sarah" as the speaker name
- If someone says "This is Mike from engineering", use "Mike" as the speaker name
- If you hear "John mentioned that he fixed the bug", use "John" as the speaker name

If you absolutely cannot identify any real names, then use Speaker_1, Speaker_2, etc.
Focus on substantial contributions, not just brief responses.

RESPONSE:"""

    try:
        response = ollama.chat(model=model_name, messages=[{"role": "user", "content": prompt}])
        content = response['message']['content'].strip()
        
        print(f"LLM Response for speaker extraction: {content[:200]}...")  # Debug output
        
        # Parse the response looking for actual names
        speakers = {}
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if ':' in line:
                # Look for patterns with actual names or speaker labels
                parts = line.split(':', 1)
                if len(parts) == 2:
                    speaker = parts[0].strip()
                    content_part = parts[1].strip()
                    
                    # Clean up speaker name (remove brackets, extra formatting)
                    speaker = speaker.replace('[', '').replace(']', '').replace('**', '').strip()
                    
                    if content_part and len(content_part) > 10:  # Filter out very short responses
                        speakers[speaker] = content_part
        
        # If parsing failed, try alternative approach
        if not speakers:
            speakers = extract_with_fallback_parsing(content, transcript)
            
        # Try to extract names from transcript directly if LLM failed
        if not speakers or all(name.startswith('Speaker_') for name in speakers.keys()):
            name_extracted_speakers = extract_names_from_transcript(transcript)
            if name_extracted_speakers:
                speakers = name_extracted_speakers
                
        # Final fallback - if still no speakers, return single speaker
        if not speakers:
            speakers = {"Speaker_1": transcript[:500] + "..." if len(transcript) > 500 else transcript}
            
        return speakers
        
    except Exception as e:
        print(f"LLM extraction failed with {model_name}: {e}")
        # Try name extraction from transcript directly
        try:
            name_speakers = extract_names_from_transcript(transcript)
            if name_speakers:
                return name_speakers
        except:
            pass
        # Ultimate fallback
        return {"Speaker_1": transcript}

def extract_names_from_transcript(transcript):
    """
    Extract speaker names directly from transcript using pattern matching
    """
    import re
    
    # Common patterns for speaker introductions
    name_patterns = [
        r"(?:hi|hello|hey),?\s+(?:i'm|i am|this is|my name is)\s+([A-Z][a-z]+)",
        r"(?:this is|here's)\s+([A-Z][a-z]+)(?:\s+(?:from|speaking|here))?",
        r"([A-Z][a-z]+)\s+(?:here|speaking|from)",
        r"(?:i'm|i am)\s+([A-Z][a-z]+)",
        r"my name is\s+([A-Z][a-z]+)",
        r"([A-Z][a-z]+)\s+mentioned",
        r"([A-Z][a-z]+)\s+said",
        r"([A-Z][a-z]+)\s+reported",
        r"([A-Z][a-z]+)\s+worked on",
        r"([A-Z][a-z]+)\s+completed",
    ]
    
    found_names = set()
    for pattern in name_patterns:
        matches = re.findall(pattern, transcript, re.IGNORECASE)
        for match in matches:
            # Filter out common words that might be mistaken for names
            if match.lower() not in ['team', 'this', 'that', 'work', 'project', 'task', 'meeting', 'update', 'status']:
                found_names.add(match.capitalize())
    
    if found_names:
        # Split transcript into segments based on sentence boundaries
        sentences = re.split(r'[.!?]+', transcript)
        speakers = {}
        
        # Assign sentences to speakers based on name proximity
        for name in found_names:
            name_content = []
            for sentence in sentences:
                sentence = sentence.strip()
                if name.lower() in sentence.lower() and len(sentence) > 20:
                    name_content.append(sentence)
            
            if name_content:
                speakers[name] = '. '.join(name_content)
            else:
                # If no specific content found, assign a portion of the transcript
                speakers[name] = f"Participated in the meeting discussion."
        
        # If we have content left over, create an "Unknown" speaker
        if len(speakers) == 0:
            return None
            
        return speakers
    
    return None

def extract_with_fallback_parsing(llm_response, original_transcript):
    """
    Alternative parsing method for LLM responses that prioritizes actual names
    """
    import re
    speakers = {}
    
    # Try to find any speaker patterns in the response, prioritizing actual names
    patterns = [
        r'([A-Z][a-z]+):\s*(.+?)(?=\n[A-Z][a-z]+:|$)',  # Actual names: content
        r'([A-Za-z_\d]+):\s*(.+?)(?=\n[A-Za-z_\d]+:|$)',  # Speaker: content
        r'([A-Z][a-z]+)\s*-\s*(.+?)(?=\n[A-Z][a-z]+\s*-|$)',  # Name - content
        r'([A-Za-z_\d]+)\s*-\s*(.+?)(?=\n[A-Za-z_\d]+\s*-|$)',  # Speaker - content
        r'([A-Z][a-z]+)\s*says?\s*:?\s*"([^"]+)"',  # Name says "content"
        r'([A-Z][a-z]+)\s*mentioned\s*:?\s*(.+?)(?=\n|$)',  # Name mentioned content
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, llm_response, re.DOTALL | re.IGNORECASE)
        for match in matches:
            if len(match) == 2:
                speaker, content = match
                speaker = speaker.strip()
                content = content.strip()
                
                # Clean up speaker name
                speaker = speaker.replace('[', '').replace(']', '').replace('**', '').strip()
                
                if content and len(content) > 20:
                    speakers[speaker] = content
        
        if speakers:
            break
    
    # If still no speakers found, try to extract names directly from original transcript
    if not speakers:
        name_based_speakers = extract_names_from_transcript(original_transcript)
        if name_based_speakers:
            speakers = name_based_speakers
        else:
            speakers = intelligent_split(original_transcript)
    
    return speakers

def intelligent_split(transcript):
    """
    Intelligently split transcript when LLM fails, attempting to identify actual names
    """
    import re
    
    # First try to extract names one more time
    name_speakers = extract_names_from_transcript(transcript)
    if name_speakers:
        return name_speakers
    
    # Look for common speaker transition patterns
    patterns = [
        r'\b(okay|so|well|um|uh|right|now),?\s+',
        r'\b(next|then|also|additionally|furthermore),?\s+',
        r'\b(hi|hello|hey)\b',
        r'\.\s+[A-Z]',  # Sentence boundaries
    ]
    
    # Split on sentence boundaries and group
    sentences = re.split(r'[.!?]+\s+', transcript)
    
    if len(sentences) <= 1:
        return {"Speaker_1": transcript}
    
    # Look for any capitalized words that might be names
    potential_names = set()
    for sentence in sentences:
        words = sentence.split()
        for word in words:
            # Look for capitalized words that aren't common words
            if (word and word[0].isupper() and len(word) > 2 and 
                word.lower() not in ['the', 'and', 'but', 'for', 'with', 'this', 'that', 'from', 'they', 'have', 'been', 'were', 'said', 'what', 'when', 'where', 'will', 'there', 'their', 'would', 'could', 'should']):
                # Simple heuristic: if it appears with pronouns or verbs, it might be a name
                if any(indicator in sentence.lower() for indicator in ['said', 'mentioned', 'worked', 'completed', 'reported', 'i am', "i'm", 'my name']):
                    potential_names.add(word)
    
    # If we found potential names, try to assign content to them
    if potential_names and len(potential_names) <= 4:  # Don't go crazy with too many names
        speakers = {}
        for name in list(potential_names)[:3]:  # Limit to 3 speakers max
            name_sentences = []
            for sentence in sentences:
                if name.lower() in sentence.lower():
                    name_sentences.append(sentence.strip())
            
            if name_sentences:
                speakers[name] = '. '.join(name_sentences)
            else:
                # Assign some content anyway
                speakers[name] = "Participated in the meeting discussion."
        
        # If we have unassigned content, add it to the first speaker
        if speakers:
            return speakers
    
    # Fallback: distribute sentences among generic speakers
    num_speakers = min(3, max(2, len(sentences) // 3))  # 2-3 speakers max
    speakers = {}
    
    sentences_per_speaker = len(sentences) // num_speakers
    
    for i in range(num_speakers):
        start_idx = i * sentences_per_speaker
        end_idx = start_idx + sentences_per_speaker if i < num_speakers - 1 else len(sentences)
        
        speaker_content = '. '.join(sentences[start_idx:end_idx])
        if speaker_content.strip():
            speakers[f"Speaker_{i + 1}"] = speaker_content.strip()
    
    return speakers if speakers else {"Speaker_1": transcript}

def extract_strict(transcript):
    import re
    updates = {}
    segments = re.split(r"Hi Team, this is (\w+)", transcript)
    for i in range(1, len(segments)-1, 2):
        name, update = segments[i], segments[i+1].strip()
        updates[name] = update
    return updates

def extract_smart(transcript):
    import spacy
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(transcript)
    updates = {}
    current_name = None
    buffer = []

    for sent in doc.sents:
        ents = [ent.text for ent in sent.ents if ent.label_ == "PERSON"]
        if ents:
            if current_name and buffer:
                updates[current_name] = " ".join(buffer).strip()
            current_name = ents[0]
            buffer = [sent.text]
        else:
            buffer.append(sent.text)
    if current_name and buffer:
        updates[current_name] = " ".join(buffer).strip()
    return updates

def summarize_updates(updates_dict, model_name):
    summaries = {}
    for name, update in updates_dict.items():
        prompt = f"Summarize this update in 1-2 sentences focusing only on work progress:\n\n{update}"
        response = ollama.chat(model=model_name, messages=[{"role": "user", "content": prompt}])
        summaries[name] = response['message']['content'].strip()
    return summaries

def format_summary_report(summary_dict, meeting_date):
    date_str = meeting_date.strftime("%Y-%m-%d")
    report_lines = [f"🗓️ Meeting Summary - {date_str}", "============================"]
    for name, summary in summary_dict.items():
        report_lines.append(f"\n👤 {name}:\n{summary}")
    return "\n".join(report_lines)

# Pyannote speaker diarization
def transcribe_with_speakers(wav_path):
    """
    Transcribe audio with speaker diarization using pyannote
    """
    try:
        from pyannote.audio import Pipeline
        import torch
        
        # Load pipeline
        pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization")
        
        # Apply diarization
        diarization = pipeline(wav_path)
        
        # Get transcript first
        transcript = transcribe_audio_faster_whisper(wav_path)
        
        # Simple speaker assignment based on time segments
        speakers = {}
        segments = list(diarization.itertracks(yield_label=True))
        
        if segments:
            # Assign speakers based on segments
            for i, (segment, _, speaker) in enumerate(segments):
                speaker_key = f"Speaker_{speaker}"
                if speaker_key not in speakers:
                    speakers[speaker_key] = []
                
                # Simple approach: divide transcript equally among speakers
                words = transcript.split()
                words_per_segment = len(words) // len(segments)
                start_idx = i * words_per_segment
                end_idx = (i + 1) * words_per_segment if i < len(segments) - 1 else len(words)
                
                segment_text = " ".join(words[start_idx:end_idx])
                speakers[speaker_key].append(segment_text)
            
            # Combine text for each speaker
            return {speaker: " ".join(texts) for speaker, texts in speakers.items()}
        else:
            return {"Speaker_1": transcript}
            
    except Exception as e:
        print(f"Pyannote diarization failed: {e}")
        # Fallback to single speaker
        transcript = transcribe_audio_faster_whisper(wav_path)
        return {"Speaker_1": transcript}

# SpeechBrain speaker recognition
def transcribe_with_speechbrain(wav_path):
    """
    Transcribe audio with speaker recognition using SpeechBrain
    """
    try:
        import torchaudio
        from speechbrain.pretrained import SpeakerRecognition
        import numpy as np
        from sklearn.cluster import AgglomerativeClustering
        
        # Load SpeechBrain model
        verification = SpeakerRecognition.from_hparams(source="speechbrain/spkrec-ecapa-voxceleb")
        
        # Load audio
        signal, fs = torchaudio.load(wav_path)
        
        # Extract speaker embeddings for different segments
        segment_length = fs * 5  # 5 second segments
        embeddings = []
        segments = []
        
        for i in range(0, signal.shape[1], segment_length):
            end = min(i + segment_length, signal.shape[1])
            segment = signal[:, i:end]
            
            if segment.shape[1] > fs:  # At least 1 second
                embedding = verification.encode_batch(segment)
                embeddings.append(embedding.squeeze().cpu().numpy())
                segments.append((i / fs, end / fs))  # Time in seconds
        
        if len(embeddings) > 1:
            # Cluster embeddings to identify speakers
            embeddings_array = np.array(embeddings)
            clustering = AgglomerativeClustering(n_clusters=min(3, len(embeddings)))
            labels = clustering.fit_predict(embeddings_array)
            
            # Get transcript
            transcript = transcribe_audio_faster_whisper(wav_path)
            words = transcript.split()
            
            # Assign words to speakers based on time segments
            speakers = {}
            for i, label in enumerate(labels):
                speaker_key = f"Speaker_{label + 1}"
                if speaker_key not in speakers:
                    speakers[speaker_key] = []
                
                # Simple word assignment
                words_per_segment = len(words) // len(labels)
                start_idx = i * words_per_segment
                end_idx = (i + 1) * words_per_segment if i < len(labels) - 1 else len(words)
                
                segment_text = " ".join(words[start_idx:end_idx])
                speakers[speaker_key].append(segment_text)
            
            return {speaker: " ".join(texts) for speaker, texts in speakers.items()}
        else:
            transcript = transcribe_audio_faster_whisper(wav_path)
            return {"Speaker_1": transcript}
            
    except Exception as e:
        print(f"SpeechBrain recognition failed: {e}")
        # Fallback to single speaker
        transcript = transcribe_audio_faster_whisper(wav_path)
        return {"Speaker_1": transcript}

# Resemblyzer speaker verification
def transcribe_with_resemblyzer(wav_path):
    """
    Transcribe audio with speaker verification using Resemblyzer
    """
    try:
        from resemblyzer import VoiceEncoder, preprocess_wav
        from sklearn.cluster import AgglomerativeClustering
        import numpy as np
        
        # Load Resemblyzer model
        encoder = VoiceEncoder()
        
        # Load and preprocess audio
        wav = preprocess_wav(wav_path)
        
        # Extract embeddings for segments
        segment_duration = 5.0  # 5 seconds
        rate = 16000  # Resemblyzer expects 16kHz
        segment_samples = int(segment_duration * rate)
        
        embeddings = []
        segments = []
        
        for i in range(0, len(wav), segment_samples):
            end = min(i + segment_samples, len(wav))
            segment = wav[i:end]
            
            if len(segment) > rate:  # At least 1 second
                embedding = encoder.embed_utterance(segment)
                embeddings.append(embedding)
                segments.append((i / rate, end / rate))
        
        if len(embeddings) > 1:
            # Cluster embeddings
            embeddings_array = np.array(embeddings)
            clustering = AgglomerativeClustering(n_clusters=min(3, len(embeddings)))
            labels = clustering.fit_predict(embeddings_array)
            
            # Get transcript
            transcript = transcribe_audio_faster_whisper(wav_path)
            words = transcript.split()
            
            # Assign words to speakers
            speakers = {}
            for i, label in enumerate(labels):
                speaker_key = f"Speaker_{label + 1}"
                if speaker_key not in speakers:
                    speakers[speaker_key] = []
                
                words_per_segment = len(words) // len(labels)
                start_idx = i * words_per_segment
                end_idx = (i + 1) * words_per_segment if i < len(labels) - 1 else len(words)
                
                segment_text = " ".join(words[start_idx:end_idx])
                speakers[speaker_key].append(segment_text)
            
            return {speaker: " ".join(texts) for speaker, texts in speakers.items()}
        else:
            transcript = transcribe_audio_faster_whisper(wav_path)
            return {"Speaker_1": transcript}
            
    except Exception as e:
        print(f"Resemblyzer verification failed: {e}")
        # Fallback to single speaker
        transcript = transcribe_audio_faster_whisper(wav_path)
        return {"Speaker_1": transcript}

# Export Functions
def export_to_word(summary_dict, meeting_date, filename=None):
    """
    Export meeting summary to Word document
    """
    if not DOCX_AVAILABLE:
        print("❌ Word export not available. Install with: pip install python-docx")
        return None
    
    try:
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date and time
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()  # Empty line
        
        # Add summary for each speaker
        for i, (speaker, summary) in enumerate(summary_dict.items()):
            # Speaker heading
            heading = doc.add_heading(f'👤 {speaker}', level=2)
            
            # Speaker content
            doc.add_paragraph(summary)
            
            # Add space between speakers (except for last one)
            if i < len(summary_dict) - 1:
                doc.add_paragraph()
        
        # Generate filename if not provided
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.docx"
        
        # Save document
        doc.save(filename)
        print(f"✅ Word document saved as: {filename}")
        return filename
        
    except Exception as e:
        print(f"❌ Error exporting to Word: {e}")
        return None

def export_to_pdf(summary_dict, meeting_date, filename=None):
    """
    Export meeting summary to PDF document
    """
    if not PDF_AVAILABLE:
        print("❌ PDF export not available. Install with: pip install reportlab")
        return None
    
    try:
        # Generate filename if not provided
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.pdf"
        
        # Create PDF document
        doc = SimpleDocTemplate(filename, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.darkblue,
            alignment=1,  # Center alignment
            spaceAfter=30
        )
        
        speaker_style = ParagraphStyle(
            'SpeakerStyle',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.darkgreen,
            spaceAfter=12
        )
        
        # Add title
        title = Paragraph(f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}', title_style)
        story.append(title)
        
        # Add generation info
        gen_info = Paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal'])
        story.append(gen_info)
        story.append(Spacer(1, 20))
        
        # Add summary for each speaker
        for speaker, summary in summary_dict.items():
            # Speaker heading
            speaker_heading = Paragraph(f'👤 {speaker}', speaker_style)
            story.append(speaker_heading)
            
            # Speaker content
            content = Paragraph(summary, styles['Normal'])
            story.append(content)
            story.append(Spacer(1, 15))
        
        # Build PDF
        doc.build(story)
        print(f"✅ PDF document saved as: {filename}")
        return filename
        
    except Exception as e:
        print(f"❌ Error exporting to PDF: {e}")
        return None

def export_to_excel(summary_dict, meeting_date, filename=None):
    """
    Export meeting summary to Excel spreadsheet
    """
    if not EXCEL_AVAILABLE:
        print("❌ Excel export not available. Install with: pip install openpyxl")
        return None
    
    try:
        # Generate filename if not provided
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.xlsx"
        
        # Create workbook and worksheet
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Meeting Summary"
        
        # Title row
        ws['A1'] = f'Meeting Summary - {meeting_date.strftime("%Y-%m-%d")}'
        ws['A1'].font = Font(size=16, bold=True, color="0066CC")
        ws.merge_cells('A1:C1')
        
        # Generation info
        ws['A2'] = f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        ws['A2'].font = Font(italic=True)
        ws.merge_cells('A2:C2')
        
        # Headers
        ws['A4'] = 'Speaker'
        ws['B4'] = 'Summary'
        ws['C4'] = 'Word Count'
        
        # Style headers
        for cell in ['A4', 'B4', 'C4']:
            ws[cell].font = Font(bold=True, color="FFFFFF")
            ws[cell].fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            ws[cell].alignment = Alignment(horizontal="center")
        
        # Add data
        row = 5
        for speaker, summary in summary_dict.items():
            ws[f'A{row}'] = speaker
            ws[f'B{row}'] = summary
            ws[f'C{row}'] = len(summary.split())
            
            # Style speaker column
            ws[f'A{row}'].font = Font(bold=True)
            ws[f'A{row}'].fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
            
            row += 1
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 80
        ws.column_dimensions['C'].width = 15
        
        # Add borders
        thin_border = openpyxl.styles.Border(
            left=openpyxl.styles.Side(style='thin'),
            right=openpyxl.styles.Side(style='thin'),
            top=openpyxl.styles.Side(style='thin'),
            bottom=openpyxl.styles.Side(style='thin')
        )
        
        for row_num in range(4, row):
            for col in ['A', 'B', 'C']:
                ws[f'{col}{row_num}'].border = thin_border
        
        # Save workbook
        wb.save(filename)
        print(f"✅ Excel document saved as: {filename}")
        return filename
        
    except Exception as e:
        print(f"❌ Error exporting to Excel: {e}")
        return None

def export_to_csv(summary_dict, meeting_date, filename=None):
    """
    Export meeting summary to CSV file
    """
    try:
        # Generate filename if not provided
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"meeting_summary_{timestamp}.csv"
        
        # Prepare data
        data = []
        
        # Add header row with metadata
        data.append(['Meeting Summary', meeting_date.strftime("%Y-%m-%d")])
        data.append(['Generated on', datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        data.append([])  # Empty row
        
        # Add column headers
        data.append(['Speaker', 'Summary', 'Word Count'])
        
        # Add speaker data
        for speaker, summary in summary_dict.items():
            word_count = len(summary.split())
            data.append([speaker, summary, word_count])
        
        # Write CSV file
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerows(data)
        
        print(f"✅ CSV file saved as: {filename}")
        return filename
        
    except Exception as e:
        print(f"❌ Error exporting to CSV: {e}")
        return None

def export_summary(summary_dict, meeting_date, formats=['word', 'pdf', 'excel', 'csv'], output_dir=None):
    """
    Export meeting summary to multiple formats
    
    Args:
        summary_dict: Dictionary with speaker names as keys and summaries as values
        meeting_date: datetime object for the meeting date
        formats: List of formats to export ['word', 'pdf', 'excel', 'csv']
        output_dir: Directory to save files (optional)
    
    Returns:
        Dictionary with format names as keys and filenames as values
    """
    results = {}
    
    # Create output directory if specified
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    # Generate base filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"meeting_summary_{timestamp}"
    
    print(f"🚀 Exporting meeting summary to {len(formats)} format(s)...")
    
    for format_type in formats:
        format_type = format_type.lower()
        
        # Generate filename with directory path if specified
        if format_type == 'word':
            filename = f"{base_name}.docx"
            if output_dir:
                filename = os.path.join(output_dir, filename)
            result = export_to_word(summary_dict, meeting_date, filename)
            
        elif format_type == 'pdf':
            filename = f"{base_name}.pdf"
            if output_dir:
                filename = os.path.join(output_dir, filename)
            result = export_to_pdf(summary_dict, meeting_date, filename)
            
        elif format_type == 'excel':
            filename = f"{base_name}.xlsx"
            if output_dir:
                filename = os.path.join(output_dir, filename)
            result = export_to_excel(summary_dict, meeting_date, filename)
            
        elif format_type == 'csv':
            filename = f"{base_name}.csv"
            if output_dir:
                filename = os.path.join(output_dir, filename)
            result = export_to_csv(summary_dict, meeting_date, filename)
            
        else:
            print(f"❌ Unknown format: {format_type}")
            result = None
        
        results[format_type] = result
    
    # Print summary
    successful_exports = [fmt for fmt, result in results.items() if result is not None]
    failed_exports = [fmt for fmt, result in results.items() if result is None]
    
    print(f"\n📊 Export Summary:")
    print(f"✅ Successful: {', '.join(successful_exports) if successful_exports else 'None'}")
    if failed_exports:
        print(f"❌ Failed: {', '.join(failed_exports)}")
    
    return results

def check_export_dependencies():
    """
    Check which export formats are available based on installed dependencies
    """
    available_formats = ['csv']  # CSV is always available
    missing_deps = []
    
    if DOCX_AVAILABLE:
        available_formats.append('word')
    else:
        missing_deps.append('python-docx (for Word export)')
    
    if PDF_AVAILABLE:
        available_formats.append('pdf')
    else:
        missing_deps.append('reportlab (for PDF export)')
    
    if EXCEL_AVAILABLE:
        available_formats.append('excel')
    else:
        missing_deps.append('openpyxl (for Excel export)')
    
    print("📋 Export Format Availability:")
    print(f"✅ Available: {', '.join(available_formats)}")
    
    if missing_deps:
        print(f"❌ Missing dependencies:")
        for dep in missing_deps:
            print(f"   - {dep}")
        print(f"\n💡 Install missing dependencies with:")
        if not DOCX_AVAILABLE:
            print("   pip install python-docx")
        if not PDF_AVAILABLE:
            print("   pip install reportlab")
        if not EXCEL_AVAILABLE:
            print("   pip install openpyxl")
    
    return available_formats
